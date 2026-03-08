"""
document_loader.py
------------------
Parses uploaded files (PDF, DOCX, TXT, HTML, images) and returns a list of
{"filename": str, "text": str} dicts that can be fed into vectorstore.add_documents().

Supported formats
-----------------
  .pdf    → PyMuPDF  (fitz)
  .docx   → python-docx
  .txt    → plain read
  .html / .htm → BeautifulSoup
  .png / .jpg / .jpeg / .bmp / .tiff / .webp → Pillow (returns a notice; no OCR)
"""

import io
import os
from pathlib import Path


# ── PDF ────────────────────────────────────────────────────────────────────────
def _load_pdf(file_bytes: bytes, filename: str) -> str:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text.strip()}")

    return "\n\n".join(text_parts)


# ── DOCX ───────────────────────────────────────────────────────────────────────
def _load_docx(file_bytes: bytes, filename: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


# ── TXT ────────────────────────────────────────────────────────────────────────
def _load_txt(file_bytes: bytes, filename: str) -> str:
    """Decode plain text, trying UTF-8 then falling back to latin-1."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


# ── HTML ───────────────────────────────────────────────────────────────────────
def _load_html(file_bytes: bytes, filename: str) -> str:
    """Extract visible text from an HTML file using BeautifulSoup."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(file_bytes, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    lines = [line.strip() for line in soup.get_text("\n").splitlines() if line.strip()]
    return "\n".join(lines)


# ── Image ──────────────────────────────────────────────────────────────────────
def _load_image(file_bytes: bytes, filename: str) -> str:
    """
    Images are displayed in the UI but we can't do OCR without extra tools.
    We return a descriptive placeholder so the chunk is still stored.
    """
    from PIL import Image

    img = Image.open(io.BytesIO(file_bytes))
    width, height = img.size
    mode = img.mode
    return (
        f"[Image file: {filename}]\n"
        f"Dimensions: {width}x{height} pixels, mode: {mode}.\n"
        "Note: Full OCR is not available. "
        "If this image contains important text, please copy it into a .txt file and upload that instead."
    )


# ── Dispatcher ────────────────────────────────────────────────────────────────
LOADERS = {
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".txt":  _load_txt,
    ".html": _load_html,
    ".htm":  _load_html,
    ".png":  _load_image,
    ".jpg":  _load_image,
    ".jpeg": _load_image,
    ".bmp":  _load_image,
    ".tiff": _load_image,
    ".webp": _load_image,
}


def load_file(file_bytes: bytes, filename: str) -> dict | None:
    """
    Parse an uploaded file and return {"filename": str, "text": str},
    or None if the file type is unsupported or parsing fails.

    Parameters
    ----------
    file_bytes : raw bytes from st.file_uploader or open(..., "rb")
    filename   : original file name (used for logging and source tracking)
    """
    ext = Path(filename).suffix.lower()
    loader = LOADERS.get(ext)

    if loader is None:
        print(f"[DocumentLoader] Unsupported file type: {ext} ({filename})")
        return None

    try:
        text = loader(file_bytes, filename)
        if not text.strip():
            print(f"[DocumentLoader] No text extracted from: {filename}")
            return None
        print(f"[DocumentLoader] Loaded {len(text):,} chars from: {filename}")
        return {"filename": filename, "text": text}
    except Exception as e:
        print(f"[DocumentLoader] Error loading {filename}: {e}")
        return None


def load_files(uploaded_files: list) -> list[dict]:
    """
    Convenience wrapper that processes a list of Streamlit UploadedFile objects.

    Returns
    -------
    list of {"filename": str, "text": str}
    """
    results = []
    for uploaded_file in uploaded_files:
        file_bytes = uploaded_file.read()
        doc = load_file(file_bytes, uploaded_file.name)
        if doc:
            results.append(doc)
    return results
